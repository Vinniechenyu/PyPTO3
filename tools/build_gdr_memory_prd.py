from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"D:\project\PyPTO3\docs\PyPTO_GDR整网接入内存优化_产品设计文档_V1.0.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "667085"
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
PALE_BLUE = "F4F8FC"
PALE_GOLD = "FFF8E6"
RED = "9B1C1C"
WHITE = "FFFFFF"

PAGE_WIDTH_DXA = 12240
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, size=None, bold=None, color=None, italic=None, ascii_font="Calibri", cjk_font="Microsoft YaHei"):
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), cjk_font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, size, color=INK, bold=False, before=0, after=6, line=1.10):
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_field(paragraph, instruction, display="1"):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, sep, text, end])
    set_run_font(run, size=9, color=MUTED)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D0D5DD", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def create_numbering(doc, kind="bullet"):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl.append(lvl_text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    ppr.append(ind)
    lvl.append(ppr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abstract_id))
    num.append(abs_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def add_list_item(doc, text, num_id, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    apply_numbering(p, num_id)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_callout(doc, label, text, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color=accent, size="8")
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(label + "  ")
    set_run_font(r, bold=True, color=accent)
    r = p.add_run(text)
    set_run_font(r, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths, alignments=None, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=font_size, bold=True, color=DARK_BLUE)
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if alignments:
                p.alignment = alignments[i]
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=INK)
        set_table_geometry(table, widths)
    return table


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    set_style_font(doc.styles["Normal"], 11, INK, False, 0, 6, 1.10)
    set_style_font(doc.styles["Heading 1"], 16, BLUE, True, 16, 8, 1.0)
    set_style_font(doc.styles["Heading 2"], 13, BLUE, True, 12, 6, 1.0)
    set_style_font(doc.styles["Heading 3"], 12, DARK_BLUE, True, 8, 4, 1.0)

    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("PyPTO 产品设计文档")
    set_run_font(r, size=9, bold=True, color=MUTED)
    r = p.add_run("\tV1.0 | 2026-08-05")
    set_run_font(r, size=9, color=MUTED)
    tabs = p.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(6.5))

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("第 ")
    set_run_font(r, size=9, color=MUTED)
    add_field(p, "PAGE")
    r = p.add_run(" 页")
    set_run_font(r, size=9, color=MUTED)


def build():
    doc = Document()
    configure_document(doc)
    bullet_num = create_numbering(doc, "bullet")
    decimal_num = create_numbering(doc, "decimal")

    # First-page memo masthead.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("产品设计文档")
    set_run_font(r, size=11, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("PyPTO 基于 GDR 算子的整网接入内存优化")
    set_run_font(r, size=24, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run("面向 Eager / ACLGraph 场景的内存治理、预算可观测与 Profiling 适配")
    set_run_font(r, size=13, color=MUTED)

    metadata = [
        ("文档版本", "V1.0（评审稿）"),
        ("需求来源", "GitCode CANN/pypto Issue #2880"),
        ("产品范围", "PyPTO Runtime、编译产物、Torch 内存池、Profiling 与调优工具"),
        ("目标周期", "Q3；需求拆解目标日期 2026-08-10"),
        ("负责人", "Issue 指定负责人：meiqing2"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(label + "：")
        set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
        r = p.add_run(value)
        set_run_font(r, size=10.5, color=INK)

    doc.add_paragraph().paragraph_format.space_after = Pt(5)
    add_callout(
        doc,
        "设计结论",
        "优先消除与算子无关的固定开销（782.90 MB 性能缓冲、100 MB 控制流缓存），再按执行模式和实际编译结果精细化 RingBuffer、Workspace 与 metadata；同时建立可归因、可对比、可验收的统一内存观测能力。",
    )

    add_heading(doc, "文档说明", 2)
    add_body(doc, "本文档基于 2026-08-05 保存的 Issue #2880 页面及其中的实测报告整理。报告中的测量值用于建立基线，不自动等同于最终产品承诺；带“待核实”的口径必须在需求评审或实现阶段闭环。")

    page_break(doc)

    add_heading(doc, "1. 执行摘要", 1)
    add_body(doc, "较大融合算子在使用 stitch_num 与 unroll 后，内部中间张量、OutCast、控制流缓存和元数据会显著放大。GDR 正向与反向用例显示，当前实现还叠加了与算子无关的进程级固定分配、固定 100 MB 控制流缓存、模式不敏感的 RingBuffer，以及难以被现有 Profiling 正确归因的 Torch/CANN 双通道内存。")
    add_body(doc, "本产品设计将内存治理拆成四条主线：固定开销门控、执行模式感知的按需分配、编译预算精细化、统一观测与调优。优先级最高的工作应先带来确定性节省，再推进 Workspace 算法和元数据估算优化。")

    add_heading(doc, "1.1 预期业务价值", 2)
    add_list_item(doc, "提升大融合算子的整网接入成功率，降低因 HBM 峰值过高导致的 OOM 与调优阻塞。", bullet_num)
    add_list_item(doc, "把“内存占用过大”转化为可定位到生命周期、申请入口、算子阶段和数据结构的明确问题。", bullet_num)
    add_list_item(doc, "让 Eager、ACLGraph、A2/A3/A5 等场景获得与实际需求匹配的分配策略，避免一刀切预留。", bullet_num)
    add_list_item(doc, "为性能调优 Agent/Skill 提供可信数据源和可执行建议，缩短人工分析链路。", bullet_num)

    add_heading(doc, "2. 背景与用户问题", 1)
    add_heading(doc, "2.1 目标用户", 2)
    add_table(
        doc,
        ["角色", "核心任务", "当前痛点"],
        [
            ("算子开发者", "开发和调优 stitch/unroll 融合算子", "无法快速判断 Workspace、OutCast、metadata 中谁主导峰值"),
            ("整网集成人员", "将 GDR 等算子接入模型并控制整网 HBM", "单算子可运行，但叠加固定开销后整网空间不足"),
            ("框架开发者", "维护 Runtime、内存池和编译预算", "Host 估算与 Device 实际使用缺少统一对账"),
            ("性能/Profiling 工程师", "采集、归因和比较内存", "atten::empty 无法区分 PyPTO；Total Allocated 不包含全部申请"),
        ],
        [1800, 3000, 4560],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )

    add_heading(doc, "2.2 核心使用场景", 2)
    for text in [
        "Eager 单算子/整网运行：没有控制流提前发射，RingBuffer 应最小化，并尽量由 Torch 内存池统一管理。",
        "ACLGraph 提前发射：保留多槽 RingBuffer，但容量应基于实际控制流数据计算。",
        "大 stitch/unroll 编译：在分配前展示 rootInner、OutCast、spill、metadata 等维度的预算与放大因子。",
        "Profiling/回归：同时统计 rtMalloc、aclrtMalloc 和 torch.empty，并能按 PyPTO 算子与生命周期聚合。",
    ]:
        add_list_item(doc, text, bullet_num)

    page_break(doc)

    add_heading(doc, "3. 现状基线与问题诊断", 1)
    add_callout(doc, "基线环境", "CANN 9.0.0，A3/910C，PYPTO_MEM_LOG=1；GDR fwd/bwd 均为 k_eff=64、parallelism=1、aicoreCount=75，测试结果 PASS。")

    add_heading(doc, "3.1 生命周期基线", 2)
    add_table(
        doc,
        ["生命周期", "主要组成", "GDR 实测/报告口径", "产品判断"],
        [
            ("进程级", "DeviceArgs、后端 SO、AICore 性能缓冲", "1.4 MB / 0.95 MB / 782.90 MB", "性能缓冲是首要固定开销，应受开关控制"),
            ("算子级", "DevAscendProgram、RingBuffer+metadata、dynamicCellMatch", "96.55-105.48 MB / 约 100 MB / 0.06-96 MB", "固定预留和按平台估算均需精细化"),
            ("每次 Launch", "Torch Workspace、用户 I/O", "382.73-519.69 MB / 322.63-589.13 MB", "需统一池化、预算解释和峰值控制"),
        ],
        [1300, 2600, 2600, 2860],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
        font_size=9,
    )
    add_callout(doc, "数据口径提醒", "Issue 的生命周期汇总写为“进程级约 885 MB”，但列项 1.4 + 0.95 + 782.90 约为 785.25 MB。可能存在未列出的约 100 MB 项或统计重复，必须通过统一采集再次核对。", fill=PALE_GOLD, accent="7A5A00")

    add_heading(doc, "3.2 正向/反向峰值对比", 2)
    add_table(
        doc,
        ["指标", "GDR fwd", "GDR bwd", "差异解释"],
        [
            ("C++ rtMalloc 小计", "980.73 MB", "1094.49 MB", "bwd 增加 dynamicCellMatch 与编译产物"),
            ("Torch Workspace", "382.73 MB", "519.69 MB", "bwd rootInner 占主导"),
            ("Torch reserved / allocated 峰值", "1040 / 898.64 MB", "3708 / 2826.95 MB", "bwd I/O 与运行期分配更高"),
            ("HBM 增量（npu-smi）", "+2223 MB", "+5024 MB", "包含框架、池化与运行时综合开销"),
        ],
        [2300, 1800, 1800, 3460],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )

    add_heading(doc, "3.3 主要根因", 2)
    roots = [
        "固定分配：非 SIM 模式下 AICore 性能缓冲无条件分配 782.90 MB，与算子无关。",
        "固定预留：DevAscendProgram 中 ctrlflowCache.cacheData 恒为 100,000,000 B（95.37 MB）。",
        "模式不敏感：Eager 不存在提前发射，但 RingBuffer 仍按多槽策略分配。",
        "放大因子：Workspace 的 rootInner、staticOutcast 受 k_eff、assemble slot、boundary/temporal slot 等影响。",
        "平台差异：A2/A3 与 A5 元数据规模不同，SlabAllocator 不能共享同一估算模型。",
        "观测割裂：CANN Runtime 与 Torch 内存池分别统计，Profiling 无法完整归因到 PyPTO 算子。",
    ]
    for root in roots:
        add_list_item(doc, root, bullet_num)

    page_break(doc)

    add_heading(doc, "4. 产品目标与边界", 1)
    add_heading(doc, "4.1 产品目标", 2)
    goals = [
        "G1：默认运行路径不为未启用的性能采集能力长期占用大块 HBM。",
        "G2：RingBuffer、CtrlflowCache、Workspace 与 metadata 按执行模式和实际编译结果定量分配。",
        "G3：在申请前后均可输出统一内存账本，实现 Host 预算与 Device 实际值对账。",
        "G4：现有 Profiling 能识别 PyPTO 内存，覆盖输入、输出、Workspace、元数据和 Runtime 直接申请。",
        "G5：将可复用的诊断与优化策略沉淀为 Agent/Skill。",
    ]
    for goal in goals:
        add_list_item(doc, goal, bullet_num)

    add_heading(doc, "4.2 非目标", 2)
    non_goals = [
        "本期不重写 Torch NPU 内存池，也不改变其通用分配与回收算法。",
        "本期不以牺牲算子正确性或提前发射能力为代价强制统一所有模式。",
        "本期不承诺直接消除全部 bwd 峰值；Workspace 算法优化需在可解释基线之上逐项推进。",
        "本期不把单一 GDR 用例的数值硬编码为所有算子或所有硬件的默认值。",
    ]
    for item in non_goals:
        add_list_item(doc, item, bullet_num)

    add_heading(doc, "4.3 设计原则", 2)
    add_body(doc, "按需优于预留，模式优于全局，测量优于猜测，生命周期优于申请点，兼容优于一次性重构。所有新增开关必须有安全默认值，所有动态缩减必须在分配前完成边界检查，并在异常时能够回退到兼容路径。")

    add_heading(doc, "5. 方案概览", 1)
    add_body(doc, "配置/执行模式 -> 编译期预算器 -> 分配计划与阈值校验 -> Runtime/Torch 统一分配标记 -> Device 实际账本 -> Profiling/Agent 建议")
    add_callout(doc, "核心策略", "先缩减确定性的固定项，再治理随算子变化的预算项；所有节省均以运行正确、性能无显著回退、可观测数据闭环为前提。")

    page_break(doc)

    add_heading(doc, "6. 功能需求设计", 1)
    add_table(
        doc,
        ["编号", "需求", "核心设计", "优先级"],
        [
            ("MEM-001", "性能缓冲门控", "AICore perf 内存按采集开关惰性创建；关闭时不分配，开启时全程复用", "P0"),
            ("MEM-002", "CtrlflowCache 精确化", "Host 可计算时按实际编码大小分配；无控制流缓存需求时不保留固定 100 MB", "P0"),
            ("MEM-003", "RingBuffer 模式感知", "Eager 默认 count=1；ACLGraph 根据提前发射深度配置并校验上限", "P0"),
            ("MEM-004", "Eager 统一池化", "非提前发射场景将 RingBuffer 与 Workspace 交由 Torch 内存池管理；复用 RTS TilingData 随路通路", "P0"),
            ("MEM-005", "内存预算报告", "按 rootInner、OutCast、spill、stitch、dynamicCellMatch、I/O 展示公式、放大因子和生命周期", "P0"),
            ("MEM-006", "硬件感知估算", "A2/A3/A5 使用独立 metadata/SlabAllocator 参数模型，并记录模型版本", "P1"),
            ("MEM-007", "Profiling 适配", "为 PyPTO 分配增加可归因标识，纳入 Total Allocated，并补齐 I/O 内存字段", "P1"),
            ("MEM-008", "调优 Agent/Skill", "依据预算和实测差异输出配置建议、热点排序与回归对比", "P1"),
        ],
        [1100, 1900, 5100, 1260],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER],
        font_size=8.8,
    )

    add_heading(doc, "6.1 MEM-001：性能缓冲门控", 2)
    add_list_item(doc, "默认关闭性能采集时，DeviceRunner::Init 不申请 76 x 10.3 MB 的性能缓冲。", bullet_num)
    add_list_item(doc, "首次启用性能采集时惰性分配；同一进程内跨算子复用，避免重复创建。", bullet_num)
    add_list_item(doc, "开关切换、并发初始化和失败回退必须线程安全；日志记录开关状态和实际分配量。", bullet_num)

    add_heading(doc, "6.2 MEM-002/003/004：执行模式感知分配", 2)
    add_list_item(doc, "Eager：RingBuffer count=1；无 Host CtrlflowCache 时固定缓存为 0；RingBuffer/Workspace 统一进入 Torch 池。", bullet_num)
    add_list_item(doc, "ACLGraph：根据提前发射深度计算 RingBuffer count；CtrlflowCache 以编译结果/Host 计算结果为准。", bullet_num)
    add_list_item(doc, "若精确计算不可用，允许回退到兼容预留值，但必须输出 fallback_reason。", bullet_num)

    page_break(doc)

    add_heading(doc, "6.3 MEM-005：内存预算与账本", 2)
    add_body(doc, "预算报告在编译完成、真正分配前生成；实际账本在分配成功后更新。两者使用同一维度命名和生命周期分类，允许直接计算 estimate_delta。")
    add_table(
        doc,
        ["字段", "说明", "示例"],
        [
            ("allocation_source", "申请入口", "rtMalloc / aclrtMalloc / torch.empty"),
            ("lifecycle", "生命周期", "process / operator / launch"),
            ("component", "内存组成", "perf / ctrlflow / rootInner / outcast / metadata / I/O"),
            ("context", "运行上下文", "device、mode、operator、phase、hardware"),
            ("bytes_estimated", "Host 预算", "519.69 MB"),
            ("bytes_requested", "实际请求", "519.69 MB"),
            ("bytes_reserved", "内存池保留", "3708 MB"),
            ("formula_factors", "放大因子", "k_eff=64、assemble_slot=…"),
            ("fallback_reason", "兼容回退原因", "host_ctrlflow_size_unavailable"),
        ],
        [2100, 3000, 4260],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
        font_size=9,
    )

    add_heading(doc, "6.4 MEM-006：硬件感知估算", 2)
    add_body(doc, "引入 hardware_profile 维度，由硬件代际选择 SlabAllocator、slot 与 dynamicCellMatch 的估算参数。参数必须版本化并可在日志中回溯；未知硬件走保守兼容策略并提示未校准。")

    add_heading(doc, "6.5 MEM-007：Profiling 适配", 2)
    add_list_item(doc, "为通过 torch.empty 产生的 PyPTO Workspace 增加稳定算子标识，避免只显示 atten::empty。", bullet_num)
    add_list_item(doc, "memory_record 的 Total Allocated 口径覆盖 PyPTO Runtime 直接申请，或提供并列的 external_runtime_allocated。", bullet_num)
    add_list_item(doc, "operator_memory 增加 Input/Output、Workspace、metadata、process-shared 字段；共享项禁止重复计入单算子总量。", bullet_num)

    add_heading(doc, "6.6 MEM-008：调优 Agent/Skill", 2)
    add_body(doc, "Skill 输入统一账本和预算报告，输出热点 Top N、Host/Device 偏差、可行动建议和预估节省。自动修改配置或代码不在默认权限内，建议需包含依据、适用模式、风险和验证用例。")

    page_break(doc)

    add_heading(doc, "7. 非功能需求", 1)
    add_table(
        doc,
        ["维度", "要求"],
        [
            ("兼容性", "默认配置保证既有用例可运行；无法精确计算时可回退至旧策略并记录原因。"),
            ("性能", "关闭诊断日志时新增判断不进入热路径；P0 优化不得造成可感知的算子性能回退。"),
            ("正确性", "所有动态容量必须完成上界、对齐、溢出和空值检查；缩减后不得发生越界或数据覆盖。"),
            ("稳定性", "进程级惰性初始化必须线程安全；OOM 时输出完整分配计划和失败项。"),
            ("可观测性", "PYPTO_MEM_LOG 保持门控；结构化字段稳定，支持前后版本对比和自动解析。"),
            ("可测试性", "每项分配策略均提供单元测试、模式矩阵测试、GDR 回归与峰值采集脚本。"),
        ],
        [1900, 7460],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )

    add_heading(doc, "8. 验收标准", 1)
    add_table(
        doc,
        ["编号", "验收项", "通过标准"],
        [
            ("AC-01", "性能缓冲关闭", "默认非采集运行中，782.90 MB perf 缓冲不产生；开启后数据采集功能正常"),
            ("AC-02", "Eager 控制流缓存", "无控制流缓存需求时，不再出现固定 100,000,000 B cacheData 预留"),
            ("AC-03", "Eager RingBuffer", "count=1，且由统一池管理；功能与性能回归通过"),
            ("AC-04", "ACLGraph 兼容", "提前发射场景按深度分配，所有控制流回归用例通过"),
            ("AC-05", "预算可解释", "fwd/bwd 报告可复现 Workspace、metadata、DevAscendProgram 的组成与公式"),
            ("AC-06", "估算对账", "同口径的 requested 与实际 Device 分配差异可量化，未归因项为 0 或明确列出"),
            ("AC-07", "Profiling 完整", "可筛选 PyPTO 算子并查看 I/O、Workspace、Runtime 直接申请和共享项"),
            ("AC-08", "回归稳定", "GDR fwd/bwd 原始用例 PASS；连续 3 次采集的关键分配量偏差不超过 2%"),
        ],
        [1100, 2400, 5860],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
        font_size=9,
    )
    add_callout(doc, "节省口径", "P0 的确定性目标是：关闭性能采集时节省 782.90 MB；无控制流缓存需求时再消除或显著缩减固定 95.37 MB cacheData。总体 HBM 降幅应按场景分别报告，不能把共享项对每个算子重复累计。")

    page_break(doc)

    add_heading(doc, "9. 交付计划与依赖", 1)
    add_table(
        doc,
        ["阶段", "目标", "主要交付物", "退出条件"],
        [
            ("Phase 0：口径冻结", "2026-08-10 前完成需求拆解", "统一内存维度、基线脚本、硬件/模式矩阵", "885 MB 口径疑点闭环；fwd/bwd 基线可复现"),
            ("Phase 1：固定项治理", "交付 MEM-001/002/003", "perf 门控、CtrlflowCache 精确化、Eager RingBuffer", "AC-01 至 AC-04 通过"),
            ("Phase 2：统一池化与预算", "交付 MEM-004/005/006", "Torch 池化、预算报告、硬件 profile", "AC-05/06 通过，A2/A3/A5 有明确策略"),
            ("Phase 3：观测与调优", "交付 MEM-007/008", "Profiling 字段与调优 Skill", "AC-07/08 通过，输出可行动建议"),
        ],
        [1900, 2200, 2900, 2360],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
        font_size=9,
    )

    add_heading(doc, "9.1 关键依赖", 2)
    for item in [
        "Issue #2873：Host 侧控制流计算能力，决定 CtrlflowCache 是否能完全按实际结果分配。",
        "Torch NPU 内存池：需要支持 PyPTO 标识透传和 RingBuffer/Workspace 统一生命周期。",
        "Profiling 工具链：operator_memory、memory_record 与 ASCEND_PROFILER_OUTPUT 字段扩展。",
        "硬件数据：A2/A3/A5 的 SlabAllocator、slot、dynamicCellMatch 样本与校准用例。",
    ]:
        add_list_item(doc, item, bullet_num)

    add_heading(doc, "10. 风险与缓解措施", 1)
    add_table(
        doc,
        ["风险", "影响", "缓解措施"],
        [
            ("精确容量低估", "越界、数据覆盖或运行失败", "保守对齐、上界检查、兼容回退、压力测试"),
            ("共享内存重复统计", "错误判断优化收益", "以生命周期为主键，共享项单列，不摊销或明确摊销规则"),
            ("模式识别错误", "Eager/ACLGraph 分配策略错配", "模式显式传递并在启动时校验；日志打印最终策略"),
            ("硬件模型漂移", "新硬件估算失准", "profile 版本化、未知硬件告警、持续校准"),
            ("Profiling 标记开销", "热路径性能回退", "关闭时零/近零开销，结构化日志异步或仅调试启用"),
            ("跨团队依赖延期", "观测闭环晚于内存优化", "P0 使用 PYPTO_MEM_LOG 独立验收，Profiling 作为 P1 并行推进"),
        ],
        [2500, 2600, 4260],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
        font_size=9,
    )

    page_break(doc)

    add_heading(doc, "11. 待决策与待核实事项", 1)
    questions = [
        "性能缓冲的真实启用条件是什么：DUMP_DEVICE_PERF、独立开关，还是运行时采集请求？",
        "生命周期汇总“约 885 MB”比列项总和多出的约 100 MB 属于哪一项？是否与 RingBuffer/metadata 重复？",
        "Eager 下 CtrlflowCache 能否完全为 0；哪些控制流元数据仍需保留？",
        "RingBuffer 纳入 Torch 池后，跨 Launch 复用与释放边界如何定义？",
        "Host 预算和 Device 实际值的允许偏差阈值是否按组件设置，而非统一百分比？",
        "A2/A3/A5 的硬件 profile 由谁维护，如何与 CANN 版本绑定？",
        "Profiling 中进程级共享项如何展示：单独页签、算子引用，还是按规则摊销？",
        "调优 Skill 的首版只做诊断建议，还是允许生成配置补丁？",
    ]
    for q in questions:
        add_list_item(doc, q, decimal_num)

    add_heading(doc, "12. 附录：关键实测数据", 1)
    add_table(
        doc,
        ["组成", "GDR fwd", "GDR bwd"],
        [
            ("tensor.Total", "363.25 MB", "517.28 MB"),
            ("rootInnerSpilledMem", "210.60 MB", "488.34 MB"),
            ("staticOutcast", "136.62 MB", "25.00 MB"),
            ("aicoreSpilled.Total", "19.49 MB", "2.34 MB"),
            ("workspaceSize", "382.73 MB", "519.69 MB"),
            ("metadata.general", "5.91 MB", "8.13 MB"),
            ("metadata.stitchPool", "19.00 MB", "19.00 MB"),
            ("metadata.dynamicCellMatch", "0.06 MB", "96.00 MB"),
            ("DevAscendProgram", "96.55 MB", "105.48 MB"),
        ],
        [4300, 2530, 2530],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
    )

    add_heading(doc, "12.1 来源", 2)
    add_body(doc, "[1] GitCode CANN/pypto Issue #2880：《[Requirement|需求建议]: 跟踪记录基于GDR算子的整网接入内存优化》，离线快照日期 2026-08-05。")
    add_body(doc, "[2] Issue 评论中的《PyPTO Device 内存申请报告（最新代码版）》，报告日期 2026-08-05。")
    add_body(doc, "[3] Issue 评论中的 Profiling 适配问题清单，包含 atten::empty 归因、Total Allocated 口径及 Input/Output 字段缺失。")

    # Keep the final source block together where possible.
    for p in doc.paragraphs[-4:]:
        p.paragraph_format.keep_together = True

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "PyPTO 基于 GDR 算子的整网接入内存优化 - 产品设计文档"
    doc.core_properties.subject = "Issue #2880 产品需求与技术方案整理"
    doc.core_properties.author = "PyPTO 产品与工程团队"
    doc.core_properties.keywords = "PyPTO,GDR,内存优化,PRD,Profiling"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
